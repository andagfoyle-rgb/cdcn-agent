"""
Web archive page and API — file explorer, upload, download, mkdir, delete.
"""
from __future__ import annotations

import json
import logging
import mimetypes
import shutil
import urllib.parse
from datetime import datetime
from pathlib import Path

from fastapi import APIRouter, File, Form, Query, Request, UploadFile
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse

from app.auth.auth import has_permission
from app.interfaces.web_auth import (
    ALLOWED_EXTS,
    MAX_UPLOAD_BYTES,
    archive_root,
    get_full_user_from_cookie,
    safe_archive_path,
)
from app.interfaces.web_templates import (
    file_icon,
    fmt_size,
    FOLDER_ICON,
    mobile_header_html,
    page_head,
    sidebar_html,
)

log = logging.getLogger(__name__)

router = APIRouter()


# ── Archive page builder ─────────────────────────────────────────────────────

def _build_archive_page(path: str, user_role: str, username: str = "", flash: str = "", flash_type: str = "info") -> str:
    """Build the archive file explorer page HTML."""
    root = archive_root()
    root.mkdir(parents=True, exist_ok=True)

    current = safe_archive_path(path)
    if current is None or not current.exists() or not current.is_dir():
        current = root
        path = ""

    # Breadcrumbs
    rel = current.relative_to(root) if current != root else Path(".")
    parts = list(rel.parts) if rel != Path(".") else []
    crumb_html = f'<a href="/archive">Archive</a>'
    acc = ""
    for p in parts:
        acc = acc + "/" + p if acc else p
        crumb_html += f' <span>/</span> <a href="/archive?path={acc}">{p}</a>'

    can_upload = has_permission(user_role, "index")
    can_delete = has_permission(user_role, "delete_documents")

    flash_html = ""
    if flash:
        flash_html = f'<div class="alert alert-{flash_type}">{flash}</div>'

    cur_path_enc = str(rel) if rel != Path(".") else ""

    toolbar_btns = ""
    if can_upload:
        toolbar_btns += f"""
        <button class="btn btn-primary btn-sm" onclick="document.getElementById('file-input').click()">
          <svg width="14" height="14" fill="none" stroke="currentColor" stroke-width="2" viewBox="0 0 24 24"><polyline points="16 16 12 12 8 16"/><line x1="12" y1="12" x2="12" y2="21"/><path d="M20.39 18.39A5 5 0 0 0 18 9h-1.26A8 8 0 1 0 3 16.3"/></svg>
          Upload Files
        </button>
        <button class="btn btn-ghost btn-sm" onclick="openMkdirModal()">
          <svg width="14" height="14" fill="none" stroke="currentColor" stroke-width="2" viewBox="0 0 24 24"><path d="M22 19a2 2 0 0 1-2 2H4a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h5l2 3h9a2 2 0 0 1 2 2z"/><line x1="12" y1="11" x2="12" y2="17"/><line x1="9" y1="14" x2="15" y2="14"/></svg>
          New Folder
        </button>
        """

    try:
        entries = sorted(current.iterdir(), key=lambda p: (p.is_file(), p.name.lower()))
    except PermissionError:
        entries = []

    rows = []
    for entry in entries:
        if entry.name.startswith("."):
            continue
        rel_path = str(entry.relative_to(root))
        enc_path = urllib.parse.quote(rel_path)
        del_btn = ""
        if can_delete:
            del_btn = f"""<button class="btn btn-danger btn-sm" onclick="confirmDelete('{enc_path}', '{entry.name}')">
              <svg width="12" height="12" fill="none" stroke="currentColor" stroke-width="2" viewBox="0 0 24 24"><polyline points="3 6 5 6 21 6"/><path d="M19 6l-1 14a2 2 0 0 1-2 2H8a2 2 0 0 1-2-2L5 6"/><path d="M10 11v6"/><path d="M14 11v6"/><path d="M9 6V4a1 1 0 0 1 1-1h4a1 1 0 0 1 1 1v2"/></svg>
            </button>"""

        if entry.is_dir():
            rows.append(f"""
            <div class="file-item">
              <div class="file-name">
                {FOLDER_ICON}
                <a href="/archive?path={enc_path}">{entry.name}</a>
              </div>
              <div class="file-size">&mdash;</div>
              <div class="file-date">&mdash;</div>
              <div class="file-actions">{del_btn}</div>
            </div>""")
        else:
            try:
                stat = entry.stat()
                size_str = fmt_size(stat.st_size)
                mtime = datetime.fromtimestamp(stat.st_mtime).strftime("%d %b %Y")
            except OSError:
                size_str = "?"
                mtime = "?"
            rows.append(f"""
            <div class="file-item">
              <div class="file-name">
                {file_icon(entry.name)}
                <a href="/api/archive/download?path={enc_path}" download="{entry.name}">{entry.name}</a>
              </div>
              <div class="file-size">{size_str}</div>
              <div class="file-date">{mtime}</div>
              <div class="file-actions">{del_btn}</div>
            </div>""")

    if not rows:
        rows.append('<div class="file-empty">This folder is empty.</div>')

    list_header = """
    <div class="file-list-header">
      <div>Name</div><div>Size</div><div>Modified</div><div></div>
    </div>"""

    upload_zone_html = ""
    if can_upload:
        upload_zone_html = f"""
        <div class="upload-zone" id="upload-zone">
          <svg width="32" height="32" fill="none" stroke="currentColor" stroke-width="1.5" viewBox="0 0 24 24"><polyline points="16 16 12 12 8 16"/><line x1="12" y1="12" x2="12" y2="21"/><path d="M20.39 18.39A5 5 0 0 0 18 9h-1.26A8 8 0 1 0 3 16.3"/></svg>
          <p>Drag &amp; drop files here, or click Upload Files above</p>
          <p class="hint">Accepted: PDF, Word, Text, Markdown, Excel, CSV, Images &nbsp;&middot;&nbsp; Max 50 MB each</p>
        </div>
        <div id="upload-progress">
          <div style="font-size:.85rem;color:var(--text-muted);">Uploading&#8230;</div>
          <div class="progress-bar-wrap"><div class="progress-bar" id="progress-bar" style="width:0%"></div></div>
        </div>
        <input type="file" id="file-input" multiple accept=".pdf,.docx,.txt,.md,.xlsx,.csv,.odt,.pptx,.jpg,.jpeg,.png,.gif" style="display:none">
        """

    mkdir_modal = f"""
    <div class="modal-backdrop" id="mkdir-modal">
      <div class="modal">
        <h3>Create New Folder</h3>
        <div class="form-group">
          <label for="folder-name">Folder name</label>
          <input type="text" id="folder-name" placeholder="e.g. Board Minutes 2025" autofocus>
        </div>
        <div class="modal-actions">
          <button class="btn btn-ghost" onclick="closeMkdirModal()">Cancel</button>
          <button class="btn btn-primary" onclick="createFolder()">Create</button>
        </div>
      </div>
    </div>
    """

    # Read sidebar JS from static file
    sidebar_js_path = Path(__file__).parent.parent / "static" / "js" / "sidebar.js"
    try:
        sidebar_js_content = sidebar_js_path.read_text()
    except FileNotFoundError:
        sidebar_js_content = ""

    archive_js = f"""
<script>
(function() {{
  var currentPath = {json.dumps(cur_path_enc)};

  // Sidebar toggle
  {sidebar_js_content}

  // ── Folder creation ───────────────────────────────────────────────────
  function openMkdirModal() {{
    document.getElementById('mkdir-modal').classList.add('open');
    document.getElementById('folder-name').focus();
  }}
  function closeMkdirModal() {{
    document.getElementById('mkdir-modal').classList.remove('open');
  }}
  window.openMkdirModal = openMkdirModal;
  window.closeMkdirModal = closeMkdirModal;

  document.getElementById('folder-name').addEventListener('keydown', function(e) {{
    if (e.key === 'Enter') createFolder();
    if (e.key === 'Escape') closeMkdirModal();
  }});

  window.createFolder = function() {{
    var name = document.getElementById('folder-name').value.trim();
    if (!name) return;
    fetch('/api/archive/mkdir', {{
      method: 'POST',
      headers: {{ 'Content-Type': 'application/json' }},
      body: JSON.stringify({{ path: currentPath, name: name }})
    }})
    .then(function(r) {{ return r.json(); }})
    .then(function(d) {{
      if (d.ok) location.reload();
      else alert(d.detail || 'Failed to create folder.');
    }})
    .catch(function() {{ alert('Request failed.'); }});
  }};

  // ── Delete ────────────────────────────────────────────────────────────
  window.confirmDelete = function(path, name) {{
    if (!confirm('Delete "' + name + '"? This cannot be undone.')) return;
    fetch('/api/archive/item?path=' + path, {{ method: 'DELETE' }})
      .then(function(r) {{ return r.json(); }})
      .then(function(d) {{
        if (d.ok) location.reload();
        else alert(d.detail || 'Delete failed.');
      }})
      .catch(function() {{ alert('Request failed.'); }});
  }};

  // ── Upload ────────────────────────────────────────────────────────────
  var fileInput = document.getElementById('file-input');
  var uploadZone = document.getElementById('upload-zone');
  var progressWrap = document.getElementById('upload-progress');
  var progressBar = document.getElementById('progress-bar');

  if (fileInput) {{
    fileInput.addEventListener('change', function() {{
      if (fileInput.files.length) uploadFiles(fileInput.files);
    }});
  }}
  if (uploadZone) {{
    uploadZone.addEventListener('dragover', function(e) {{
      e.preventDefault();
      uploadZone.classList.add('drag-over');
    }});
    uploadZone.addEventListener('dragleave', function() {{
      uploadZone.classList.remove('drag-over');
    }});
    uploadZone.addEventListener('drop', function(e) {{
      e.preventDefault();
      uploadZone.classList.remove('drag-over');
      if (e.dataTransfer.files.length) uploadFiles(e.dataTransfer.files);
    }});
  }}

  function uploadFiles(files) {{
    var formData = new FormData();
    for (var i = 0; i < files.length; i++) formData.append('files', files[i]);
    formData.append('path', currentPath);

    progressWrap.style.display = 'block';
    progressBar.style.width = '0%';

    var xhr = new XMLHttpRequest();
    xhr.open('POST', '/api/archive/upload');
    xhr.upload.onprogress = function(e) {{
      if (e.lengthComputable) {{
        progressBar.style.width = Math.round(e.loaded / e.total * 100) + '%';
      }}
    }};
    xhr.onload = function() {{
      progressWrap.style.display = 'none';
      var d;
      try {{ d = JSON.parse(xhr.responseText); }} catch(e) {{ d = {{}}; }}
      if (xhr.status === 200 && d.ok) {{
        location.reload();
      }} else {{
        alert(d.detail || 'Upload failed.');
      }}
    }};
    xhr.onerror = function() {{
      progressWrap.style.display = 'none';
      alert('Upload request failed.');
    }};
    xhr.send(formData);
  }}
}}());
</script>
"""

    return f"""<!DOCTYPE html>
<html lang="en">
{page_head("Archive — CDCN Agent")}
<body>
  <div class="app">
    {sidebar_html("archive", username=username, role=user_role)}
    <div class="main-content">
      {mobile_header_html("Archive")}
      <div class="main-topbar">
        <span class="main-topbar-title">Document Archive</span>
        <button class="icon-btn" id="theme-toggle" onclick="toggleTheme()" title="Toggle theme"></button>
      </div>
      <div class="page-content">
        {flash_html}
        <div class="breadcrumb">{crumb_html}</div>
        <div class="explorer-toolbar">{toolbar_btns}</div>
        {upload_zone_html}
        <div class="file-list">
          {list_header}
          {"".join(rows)}
        </div>
      </div>
    </div>
  </div>
  {mkdir_modal}
  {archive_js}
</body>
</html>"""


@router.get("/archive", response_class=HTMLResponse)
async def archive_page(request: Request, path: str = Query(default="")):
    user = get_full_user_from_cookie(request)
    if not user:
        return RedirectResponse(url="/login", status_code=303)
    html = _build_archive_page(path, user.role, username=user.username)
    return HTMLResponse(html)


# ── Archive API endpoints ────────────────────────────────────────────────────

@router.get("/api/archive/ls")
async def archive_ls(request: Request, path: str = Query(default="")):
    user = get_full_user_from_cookie(request)
    if not user:
        return JSONResponse({"detail": "Unauthorised"}, status_code=401)
    target = safe_archive_path(path)
    if target is None or not target.exists() or not target.is_dir():
        return JSONResponse({"detail": "Invalid path"}, status_code=400)
    root = archive_root()
    items = []
    for entry in sorted(target.iterdir(), key=lambda p: (p.is_file(), p.name.lower())):
        if entry.name.startswith("."):
            continue
        rel = str(entry.relative_to(root))
        item: dict = {"name": entry.name, "path": rel}
        if entry.is_dir():
            item["type"] = "folder"
        else:
            try:
                s = entry.stat()
                item["type"] = "file"
                item["size"] = s.st_size
                item["modified"] = datetime.fromtimestamp(s.st_mtime).isoformat()
            except OSError:
                item["type"] = "file"
        items.append(item)
    return JSONResponse({"ok": True, "items": items})


@router.get("/api/archive/download")
async def archive_download(request: Request, path: str = Query(...)):
    user = get_full_user_from_cookie(request)
    if not user:
        return JSONResponse({"detail": "Unauthorised"}, status_code=401)
    target = safe_archive_path(path)
    if target is None or not target.exists() or not target.is_file():
        return JSONResponse({"detail": "File not found"}, status_code=404)
    mime, _ = mimetypes.guess_type(str(target))
    mime = mime or "application/octet-stream"
    from starlette.responses import FileResponse
    return FileResponse(path=str(target), media_type=mime, filename=target.name)


async def _detect_document_refinement(
    original_path: Path, new_content: bytes, username: str
) -> None:
    """Compare uploaded file with existing original to detect refinements."""
    try:
        ext = original_path.suffix.lower()
        if ext not in (".txt", ".md", ".docx"):
            return
        if ext == ".docx":
            try:
                from docx import Document as DocxDocument
                orig_doc = DocxDocument(str(original_path))
                orig_text = "\n".join(p.text for p in orig_doc.paragraphs)
                import io
                new_doc = DocxDocument(io.BytesIO(new_content))
                new_text = "\n".join(p.text for p in new_doc.paragraphs)
            except Exception:
                return
        else:
            orig_text = original_path.read_text(errors="replace")
            new_text = new_content.decode("utf-8", errors="replace")

        if not orig_text.strip() or not new_text.strip():
            return

        import difflib
        diff_lines = list(difflib.unified_diff(
            orig_text.splitlines(), new_text.splitlines(),
            fromfile="original", tofile="corrected",
            lineterm="", n=2,
        ))
        if len(diff_lines) < 3:
            return
        diff_text = "\n".join(diff_lines[:80])
        from app.storage.audit_log import log_learned_skill
        await log_learned_skill(
            skill_name="document_refinement",
            trigger_pattern=f"Drafting {ext} document: {original_path.name}",
            description=(
                f"User '{username}' uploaded a corrected version of "
                f"'{original_path.name}'. Key changes:\n{diff_text}"
            ),
            source=f"upload_refinement_{datetime.now().strftime('%Y-%m-%d')}",
        )
        log.info("Document refinement detected: %s (by %s)", original_path.name, username)
    except Exception as exc:
        log.warning("_detect_document_refinement failed: %s", exc)


@router.post("/api/archive/upload")
async def archive_upload(
    request: Request,
    files: list[UploadFile] = File(...),
    path: str = Form(default=""),
):
    user = get_full_user_from_cookie(request)
    if not user:
        return JSONResponse({"detail": "Unauthorised"}, status_code=401)
    if not has_permission(user.role, "index"):
        return JSONResponse({"detail": "Permission denied"}, status_code=403)
    target_dir = safe_archive_path(path)
    if target_dir is None or not target_dir.exists() or not target_dir.is_dir():
        return JSONResponse({"detail": "Invalid destination path"}, status_code=400)

    saved = []
    errors = []
    for upload in files:
        name = Path(upload.filename or "upload").name
        ext = Path(name).suffix.lower()
        if ext not in ALLOWED_EXTS:
            errors.append(f"{name}: file type not allowed")
            continue
        dest = target_dir / name
        existing_original = dest if dest.exists() else None
        counter = 1
        while dest.exists():
            stem = Path(name).stem
            dest = target_dir / f"{stem}_{counter}{ext}"
            counter += 1
        try:
            content = await upload.read()
            if len(content) > MAX_UPLOAD_BYTES:
                errors.append(f"{name}: file too large (max 10 MB)")
                continue
            try:
                import magic
                mime = magic.from_buffer(content, mime=True)
                _BLOCKED_MIMES = {
                    "application/x-executable", "application/x-sharedlib",
                    "application/x-mach-binary", "application/x-dosexec",
                    "application/x-shellscript", "text/x-shellscript",
                }
                if mime in _BLOCKED_MIMES:
                    errors.append(f"{name}: blocked file type detected")
                    continue
            except ImportError:
                pass
            if existing_original:
                await _detect_document_refinement(existing_original, content, user.username)
            dest.write_bytes(content)
            saved.append(dest.name)
            log.info("Archive upload: user=%s saved %s", user.username, dest)
        except Exception as exc:
            log.error("Archive upload error: %s", exc)
            errors.append(f"{name}: write failed")

    if not saved and errors:
        return JSONResponse({"ok": False, "detail": "; ".join(errors)}, status_code=400)
    return JSONResponse({"ok": True, "saved": saved, "errors": errors})


@router.post("/api/archive/mkdir")
async def archive_mkdir(request: Request):
    user = get_full_user_from_cookie(request)
    if not user:
        return JSONResponse({"detail": "Unauthorised"}, status_code=401)
    if not has_permission(user.role, "index"):
        return JSONResponse({"detail": "Permission denied"}, status_code=403)
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"detail": "Invalid JSON"}, status_code=400)
    parent_path = str(body.get("path", "") or "")
    folder_name = str(body.get("name", "")).strip()
    if not folder_name or "/" in folder_name or "\\" in folder_name or folder_name in (".", ".."):
        return JSONResponse({"detail": "Invalid folder name"}, status_code=400)
    parent = safe_archive_path(parent_path)
    if parent is None or not parent.exists() or not parent.is_dir():
        return JSONResponse({"detail": "Invalid parent path"}, status_code=400)
    new_dir = parent / folder_name
    if new_dir.exists():
        return JSONResponse({"detail": "A folder with that name already exists"}, status_code=409)
    try:
        new_dir.mkdir(parents=False)
        log.info("Archive mkdir: user=%s created %s", user.username, new_dir)
        return JSONResponse({"ok": True})
    except Exception as exc:
        log.error("Archive mkdir error: %s", exc)
        return JSONResponse({"detail": "Failed to create folder"}, status_code=500)


@router.delete("/api/archive/item")
async def archive_delete(request: Request, path: str = Query(...)):
    user = get_full_user_from_cookie(request)
    if not user:
        return JSONResponse({"detail": "Unauthorised"}, status_code=401)
    if not has_permission(user.role, "delete_documents"):
        return JSONResponse({"detail": "Permission denied"}, status_code=403)
    target = safe_archive_path(path)
    root = archive_root()
    if target is None or not target.exists():
        return JSONResponse({"detail": "Path not found"}, status_code=404)
    if target == root:
        return JSONResponse({"detail": "Cannot delete the archive root"}, status_code=400)
    try:
        if target.is_dir():
            shutil.rmtree(target)
        else:
            target.unlink()
        log.info("Archive delete: user=%s deleted %s", user.username, target)
        return JSONResponse({"ok": True})
    except Exception as exc:
        log.error("Archive delete error: %s", exc)
        return JSONResponse({"detail": "Delete failed"}, status_code=500)
