"""
DocumentEditorSkill — read, edit, and create documents in the CDCN archive.

Supports:
  • Plain text files (.txt, .md, .json, .csv, .yml, .yaml, .html, .xml)
  • Word documents (.docx) — read text content, create/overwrite via python-docx
  • PDF files (.pdf) — read-only via pdfplumber

Unlike search_archive (which returns fragments), this skill gives full document
content and lets the agent save modifications or new documents.
"""
from __future__ import annotations

import logging
import re
import shutil
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from app.config import settings
from app.skills.base import BaseSkill, SkillResult

log = logging.getLogger(__name__)

_MAX_VERSIONS = 5

# Extensions that can be read as UTF-8 text
_TEXT_EXTS = {".txt", ".md", ".json", ".csv", ".yml", ".yaml", ".html", ".xml"}

# Extensions with special handling
_DOCX_EXTS = {".docx"}
_PDF_EXTS = {".pdf"}

_ALL_READ_EXTS = _TEXT_EXTS | _DOCX_EXTS | _PDF_EXTS
_ALL_WRITE_EXTS = _TEXT_EXTS | _DOCX_EXTS  # PDFs are read-only


# ── Helpers ──────────────────────────────────────────────────────────────────

def _base_dir() -> Path:
    """Resolved archive root — all paths are restricted to this tree."""
    return Path(settings.watched_folder).resolve()


def _resolve_path(rel: str) -> tuple[Path | None, str | None]:
    """Validate *rel* is inside the archive.  Returns (abs_path, error)."""
    if not rel:
        return None, "path is required."
    base = _base_dir()
    full = (base / rel).resolve()
    try:
        full.relative_to(base)
    except ValueError:
        return None, f"Access denied: path must be within the document archive."
    return full, None


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _read_docx(path: Path) -> str:
    """Extract plain text from a .docx file, preserving paragraph breaks."""
    from docx import Document
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs]
    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def _read_pdf(path: Path) -> str:
    """Extract text from a PDF using pdfplumber with OCR fallback for scanned pages."""
    import pdfplumber
    pages: list[str] = []
    page_count = 0
    with pdfplumber.open(path) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text)
    # If average chars per page < 100, assume scanned and try OCR
    total_chars = sum(len(p) for p in pages)
    if total_chars / max(page_count, 1) < 100:
        try:
            import pytesseract
            ocr_pages: list[str] = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    img = page.to_image(resolution=300).original
                    text = pytesseract.image_to_string(img)
                    if text.strip():
                        ocr_pages.append(text.strip())
            if ocr_pages:
                pages = ocr_pages
        except ImportError:
            log.warning("OCR fallback unavailable (pytesseract not installed) for %s", path.name)
        except Exception as exc:
            log.warning("OCR fallback failed for %s: %s", path.name, exc)
    return "\n\n".join(pages)


def _write_text(path: Path, content: str) -> None:
    path.write_text(content, encoding="utf-8")


def _write_docx(path: Path, content: str, title: str = "") -> None:
    """Create a simple .docx from plain text or markdown-ish content.

    For branded CDCN documents the agent should use the writer skill +
    docx_converter instead.  This is for quick edits and simple documents.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Style defaults
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    if title:
        doc.add_heading(title, level=1)

    for line in content.split("\n"):
        stripped = line.strip()
        # Detect markdown-style headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s*", "", stripped), style="List Number")
        else:
            doc.add_paragraph(line)

    doc.save(str(path))


def _file_meta(path: Path, content: str) -> dict[str, Any]:
    stat = path.stat()
    return {
        "size_bytes": stat.st_size,
        "line_count": content.count("\n") + 1 if content else 0,
        "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
    }


def _versions_dir(path: Path) -> Path:
    """Return the .versions/ subdirectory for *path*'s parent."""
    return path.parent / ".versions"


def _version_prefix(path: Path) -> str:
    """Return the stem used to match versions of *path*."""
    return path.name


def _create_version(path: Path) -> Path | None:
    """Copy *path* into .versions/ with a timestamp.  Prune to _MAX_VERSIONS."""
    if not path.exists():
        return None
    vdir = _versions_dir(path)
    vdir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H-%M-%S")
    versioned_name = f"{path.stem}.{ts}{path.suffix}"
    dest = vdir / versioned_name
    shutil.copy2(path, dest)
    # Prune old versions — keep newest _MAX_VERSIONS
    _prune_versions(path)
    return dest


def _list_versions(path: Path) -> list[dict[str, Any]]:
    """Return metadata for all versions of *path*, newest first."""
    vdir = _versions_dir(path)
    if not vdir.is_dir():
        return []
    prefix = path.stem + "."
    suffix = path.suffix
    versions = []
    for f in sorted(vdir.iterdir(), reverse=True):
        if f.name.startswith(prefix) and f.name.endswith(suffix) and f.is_file():
            # Extract timestamp from name: stem.YYYY-MM-DDTHH-MM-SS.ext
            ts_part = f.name[len(prefix):-len(suffix)] if suffix else f.name[len(prefix):]
            stat = f.stat()
            versions.append({
                "filename": f.name,
                "timestamp": ts_part.replace("T", " ").replace("-", ":", 2) if "T" in ts_part else ts_part,
                "size_bytes": stat.st_size,
                "path": str(f),
            })
    return versions


def _prune_versions(path: Path) -> None:
    """Keep only the newest _MAX_VERSIONS versions of *path*."""
    versions = _list_versions(path)
    for old in versions[_MAX_VERSIONS:]:
        try:
            Path(old["path"]).unlink()
        except OSError:
            pass


# ── Skill class ──────────────────────────────────────────────────────────────

class DocumentEditorSkill(BaseSkill):
    name = "document_editor"
    description = (
        "Open, read, edit, and save documents in the CDCN archive. "
        "Supports .docx, .pdf (read-only), and common text formats."
    )

    async def run(self, **kwargs) -> SkillResult:
        action = kwargs.get("action", "read")
        if action == "read":
            return await self._read(kwargs)
        elif action == "save":
            return await self._save(kwargs)
        elif action == "list":
            return await self._list(kwargs)
        elif action == "list_versions":
            return await self._list_versions(kwargs)
        elif action == "restore_version":
            return await self._restore_version(kwargs)
        else:
            return SkillResult(
                success=False,
                error=f"Unknown action '{action}'.  Use 'read', 'save', 'list', 'list_versions', or 'restore_version'.",
            )

    # ── read ─────────────────────────────────────────────────────────────

    async def _read(self, kw: dict) -> SkillResult:
        rel = kw.get("path") or kw.get("document_path")
        if not rel:
            return SkillResult(success=False, error="'path' is required.")

        full, err = _resolve_path(rel)
        if err:
            return SkillResult(success=False, error=err)
        if not full.exists():
            return SkillResult(success=False, error=f"Not found: {rel}")
        if not full.is_file():
            return SkillResult(success=False, error=f"Not a file: {rel}")

        ext = full.suffix.lower()
        if ext not in _ALL_READ_EXTS:
            return SkillResult(
                success=False,
                error=f"Unsupported file type '{ext}'.  Supported: {', '.join(sorted(_ALL_READ_EXTS))}",
            )

        try:
            if ext in _PDF_EXTS:
                content = _read_pdf(full)
            elif ext in _DOCX_EXTS:
                content = _read_docx(full)
            else:
                content = _read_text(full)

            meta = _file_meta(full, content)
            return SkillResult(
                success=True,
                output={
                    "path": rel,
                    "extension": ext,
                    "content": content,
                    **meta,
                },
            )
        except Exception as exc:
            return SkillResult(success=False, error=f"Error reading {rel}: {exc}")

    # ── save ─────────────────────────────────────────────────────────────

    async def _save(self, kw: dict) -> SkillResult:
        rel = kw.get("path") or kw.get("output_path") or kw.get("document_path")
        content = kw.get("content")
        create_backup = kw.get("create_backup", True)
        title = kw.get("title", "")

        if content is None:
            return SkillResult(success=False, error="'content' is required for save.")
        if not rel:
            return SkillResult(success=False, error="'path' is required for save.")

        full, err = _resolve_path(rel)
        if err:
            return SkillResult(success=False, error=err)

        ext = full.suffix.lower()
        if ext not in _ALL_WRITE_EXTS:
            return SkillResult(
                success=False,
                error=f"Cannot write '{ext}' files.  Writable: {', '.join(sorted(_ALL_WRITE_EXTS))}",
            )

        try:
            version_path = None
            if create_backup:
                version_path = _create_version(full)

            full.parent.mkdir(parents=True, exist_ok=True)

            if ext in _DOCX_EXTS:
                _write_docx(full, content, title=title)
            else:
                _write_text(full, content)

            result: dict[str, Any] = {
                "path": rel,
                "size_bytes": full.stat().st_size,
                "action": "saved",
            }
            if version_path:
                result["version_saved"] = version_path.name
            return SkillResult(success=True, output=result)

        except Exception as exc:
            return SkillResult(success=False, error=f"Error saving {rel}: {exc}")

    # ── list ─────────────────────────────────────────────────────────────

    async def _list(self, kw: dict) -> SkillResult:
        """List files in a directory, optionally filtered by extension."""
        rel = kw.get("path", "")
        ext_filter = kw.get("extension")  # e.g. ".docx"

        base = _base_dir()
        target = (base / rel).resolve() if rel else base
        try:
            target.relative_to(base)
        except ValueError:
            return SkillResult(success=False, error="Access denied.")

        if not target.is_dir():
            return SkillResult(success=False, error=f"Not a directory: {rel}")

        entries = []
        for item in sorted(target.iterdir()):
            if item.name.startswith("."):
                continue
            if item.is_dir():
                entries.append({"name": item.name, "type": "directory"})
            elif item.is_file():
                if ext_filter and item.suffix.lower() != ext_filter.lower():
                    continue
                entries.append({
                    "name": item.name,
                    "type": "file",
                    "extension": item.suffix.lower(),
                    "size_bytes": item.stat().st_size,
                })

        return SkillResult(success=True, output={"path": rel or "/", "entries": entries})

    # ── list_versions ────────────────────────────────────────────────────

    async def _list_versions(self, kw: dict) -> SkillResult:
        """List saved versions of a document."""
        rel = kw.get("path") or kw.get("document_path")
        if not rel:
            return SkillResult(success=False, error="'path' is required.")

        full, err = _resolve_path(rel)
        if err:
            return SkillResult(success=False, error=err)
        if not full.exists():
            return SkillResult(success=False, error=f"Not found: {rel}")

        versions = _list_versions(full)
        if not versions:
            return SkillResult(
                success=True,
                output=f"No previous versions found for {rel}.",
                metadata={"versions": 0},
            )

        lines = [f"**{len(versions)} version(s) of {rel}:**\n"]
        for i, v in enumerate(versions, 1):
            size_kb = v["size_bytes"] / 1024
            lines.append(f"{i}. `{v['filename']}` — {v['timestamp']} ({size_kb:.1f} KB)")

        return SkillResult(
            success=True,
            output="\n".join(lines),
            metadata={"versions": len(versions), "files": [v["filename"] for v in versions]},
        )

    # ── restore_version ──────────────────────────────────────────────────

    async def _restore_version(self, kw: dict) -> SkillResult:
        """Restore a previous version of a document."""
        rel = kw.get("path") or kw.get("document_path")
        version_filename = kw.get("version_filename")

        if not rel:
            return SkillResult(success=False, error="'path' is required.")
        if not version_filename:
            return SkillResult(success=False, error="'version_filename' is required.")

        full, err = _resolve_path(rel)
        if err:
            return SkillResult(success=False, error=err)

        vdir = _versions_dir(full)
        version_path = (vdir / version_filename).resolve()
        # Security: ensure version_path is inside .versions dir
        try:
            version_path.relative_to(vdir.resolve())
        except ValueError:
            return SkillResult(success=False, error="Access denied.")
        if not version_path.exists():
            return SkillResult(success=False, error=f"Version not found: {version_filename}")

        # Save current as a version before restoring
        if full.exists():
            _create_version(full)

        shutil.copy2(version_path, full)
        log.info("Restored %s from version %s", rel, version_filename)
        return SkillResult(
            success=True,
            output=f"Restored `{rel}` from version `{version_filename}`.",
            metadata={"restored_from": version_filename, "path": rel},
        )
