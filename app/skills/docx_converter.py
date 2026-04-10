"""
Markdown -> DOCX converter for CDCN Agent.

Converts markdown text to a professionally formatted Word document.
Primary path: parses markdown into a JSON structure and delegates to
docx_generator.js (Node.js, using the `docx` npm package) for superior
formatting with full CDCN branding — title page, Rubik font, branded
tables, blockquote callouts, resolution/action boxes, page breaks, etc.

Falls back to the legacy python-docx renderer if the JS generator is
unavailable or fails.

Brand guidelines
----------------
  * Font family:  Rubik (fallback: Calibri)
  * Heading text:  #2C3E50 (Charcoal Navy)
  * Heading 2:     blue underline in Sailing Blue #5EAFE5
  * Body text:     #333333
  * Table headers:  #2C3E50 background, white text
  * Table alt rows: #F2F5F8
  * Accent colour:  Sailing Blue #5EAFE5
"""
import re
import json
import logging
import subprocess
from pathlib import Path
from datetime import datetime

log = logging.getLogger(__name__)

OUTPUT_DIR = Path("/var/lib/cdcn-agent/documents")
DOWNLOAD_BASE = "/api/archive/download?path="
JS_GENERATOR = Path(__file__).with_name("docx_generator.js")

# Regex for **Key:** Value metadata lines
_METADATA_RE = re.compile(r"^\*\*(.+?):\*\*\s*(.+)$")


def _is_metadata_line(line: str) -> bool:
    return bool(_METADATA_RE.match(line.strip()))


# ── Markdown table parser ────────────────────────────────────────────────────

def _parse_md_table(lines: list[str]) -> dict:
    """Parse markdown table lines into a JSON table block."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        if re.match(r"^\|[\s\-:|]+\|$", line):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if cells:
            rows.append(cells)
    if not rows:
        return None
    return {
        "type": "table",
        "headers": rows[0],
        "rows": rows[1:] if len(rows) > 1 else [],
    }


# ── Markdown → JSON structure ────────────────────────────────────────────────

def _markdown_to_json(md_text: str, title: str = "") -> dict:
    """
    Parse markdown text into the JSON structure expected by docx_generator.js.

    Returns a dict with keys: title, metadata, sections.
    """
    lines = md_text.split("\n")
    doc_title = title
    metadata = []
    sections = []
    current_section = None

    i = 0
    in_table = False
    table_lines: list[str] = []
    in_blockquote = False
    blockquote_lines: list[str] = []
    first_heading_seen = False
    collecting_metadata = False

    def flush_table():
        nonlocal in_table, table_lines
        if in_table and table_lines:
            tbl = _parse_md_table(table_lines)
            if tbl and current_section is not None:
                current_section["body"].append(tbl)
        in_table = False
        table_lines = []

    def flush_blockquote():
        nonlocal in_blockquote, blockquote_lines
        if in_blockquote and blockquote_lines:
            for bq_line in blockquote_lines:
                if current_section is not None:
                    # Detect resolution/action lines
                    upper = bq_line.upper()
                    if upper.startswith("ACTION:") or upper.startswith("RESOLUTION:") or upper.startswith("AGREED:") or upper.startswith("DECISION:"):
                        current_section["body"].append({
                            "type": "resolution",
                            "text": bq_line,
                        })
                    else:
                        current_section["body"].append({
                            "type": "blockquote",
                            "text": bq_line,
                        })
        in_blockquote = False
        blockquote_lines = []

    def flush_metadata():
        nonlocal collecting_metadata
        if collecting_metadata and metadata:
            pass  # metadata already collected at doc level
        collecting_metadata = False

    def ensure_section():
        nonlocal current_section
        if current_section is None:
            current_section = {"heading": "", "level": 2, "body": []}
            sections.append(current_section)

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip HTML comments
        if stripped.startswith("<!--"):
            while "-->" not in stripped and i < len(lines) - 1:
                i += 1
                stripped = lines[i].strip()
            i += 1
            continue

        # Flush blockquote if we've left it
        if in_blockquote and not stripped.startswith(">"):
            flush_blockquote()

        # Flush table if we've left it
        if in_table and not (stripped.startswith("|") and "|" in stripped[1:]):
            flush_table()

        # Flush metadata if we've left it
        if collecting_metadata and not _is_metadata_line(stripped) and stripped:
            flush_metadata()

        # ── Horizontal rule ──────────────────────────────────────────────
        if re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
            ensure_section()
            current_section["body"].append({"type": "divider"})
            i += 1
            continue

        # ── Table detection ──────────────────────────────────────────────
        if stripped.startswith("|") and "|" in stripped[1:]:
            ensure_section()
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(stripped)
            i += 1
            continue

        # ── Blockquote ───────────────────────────────────────────────────
        if stripped.startswith(">"):
            ensure_section()
            in_blockquote = True
            bq_text = stripped[1:].strip()
            if bq_text:
                blockquote_lines.append(bq_text)
            i += 1
            continue

        # ── Headings ─────────────────────────────────────────────────────
        heading_match = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if heading_match:
            # Flush any pending content
            flush_table()
            flush_blockquote()

            level = min(len(heading_match.group(1)), 3)
            heading_text = heading_match.group(2).strip()

            if not first_heading_seen:
                first_heading_seen = True
                if not doc_title:
                    doc_title = heading_text

            current_section = {
                "heading": heading_text,
                "level": level,
                "body": [],
            }
            sections.append(current_section)
            i += 1
            continue

        # ── Metadata lines (**Key:** Value) before first heading ─────────
        if _is_metadata_line(stripped) and not first_heading_seen:
            collecting_metadata = True
            m = _METADATA_RE.match(stripped)
            if m:
                metadata.append({"key": m.group(1), "value": m.group(2)})
            i += 1
            continue

        # ── Metadata lines within body ───────────────────────────────────
        if _is_metadata_line(stripped) and first_heading_seen:
            ensure_section()
            m = _METADATA_RE.match(stripped)
            if m:
                # Detect resolution/action metadata
                key_upper = m.group(1).upper()
                if key_upper in ("ACTION", "RESOLUTION", "AGREED", "DECISION"):
                    current_section["body"].append({
                        "type": "resolution",
                        "text": f"{m.group(1)}: {m.group(2)}",
                    })
                else:
                    current_section["body"].append({
                        "type": "metadata_line",
                        "key": m.group(1),
                        "value": m.group(2),
                    })
            i += 1
            continue

        # ── Sub-bullet lists (indented) ──────────────────────────────────
        indent_match = re.match(r"^(\s{2,})([-*])\s+(.*)", line)
        if indent_match:
            ensure_section()
            bullet_text = indent_match.group(3)
            # Append to the last bullets block if possible
            body = current_section["body"]
            if body and body[-1]["type"] == "bullets":
                body[-1]["items"].append("  " + bullet_text)
            else:
                current_section["body"].append({
                    "type": "bullets",
                    "items": ["  " + bullet_text],
                })
            i += 1
            continue

        # ── Bullet points ────────────────────────────────────────────────
        if stripped.startswith("- ") or stripped.startswith("* "):
            ensure_section()
            bullet_text = stripped[2:]
            body = current_section["body"]
            if body and body[-1]["type"] == "bullets":
                body[-1]["items"].append(bullet_text)
            else:
                current_section["body"].append({
                    "type": "bullets",
                    "items": [bullet_text],
                })
            i += 1
            continue

        # ── Numbered lists ───────────────────────────────────────────────
        num_match = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if num_match:
            ensure_section()
            body = current_section["body"]
            if body and body[-1]["type"] == "numbered":
                body[-1]["items"].append(num_match.group(2))
            else:
                current_section["body"].append({
                    "type": "numbered",
                    "items": [num_match.group(2)],
                })
            i += 1
            continue

        # ── Empty line ───────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Regular paragraph ────────────────────────────────────────────
        ensure_section()
        current_section["body"].append({
            "type": "paragraph",
            "text": stripped,
        })
        i += 1

    # Flush remaining
    flush_table()
    flush_blockquote()

    return {
        "title": doc_title or "CDCN Document",
        "metadata": metadata,
        "sections": sections,
    }


# ── JS generator call ───────────────────────────────────────────────────────

def _call_js_generator(json_payload: dict, filename: str) -> Path:
    """
    Call docx_generator.js via subprocess, passing JSON on stdin.
    Returns the Path to the generated DOCX.
    """
    json_payload["filename"] = filename
    json_payload["outputDir"] = str(OUTPUT_DIR)

    input_data = json.dumps(json_payload, ensure_ascii=False)

    result = subprocess.run(
        ["node", str(JS_GENERATOR)],
        input=input_data,
        capture_output=True,
        text=True,
        timeout=30,
        cwd=str(JS_GENERATOR.parent.parent.parent),  # /opt/cdcn-agent
    )

    if result.returncode != 0:
        raise RuntimeError(
            f"docx_generator.js failed (rc={result.returncode}): {result.stderr.strip()}"
        )

    out_path = Path(result.stdout.strip())
    if not out_path.exists():
        raise RuntimeError(f"JS generator reported path does not exist: {out_path}")

    return out_path


# ── Legacy python-docx fallback ──────────────────────────────────────────────

def _legacy_markdown_to_docx(md_text: str, filename: str) -> Path:
    """
    Original python-docx based converter — used as fallback when the JS
    generator is unavailable.
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    from app.skills.docx_converter_helpers import (
        CHARCOAL, SUBTLE_GREY, BODY_TEXT, DIVIDER_HEX,
        _set_run_font, _setup_styles, _add_header_footer,
        _add_h2_underline, _add_page_break, _add_rich_text,
        _add_table, _METADATA_RE, _is_metadata_line,
        _add_metadata_block, _flush_blockquote,
    )

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    _setup_styles(doc)
    _add_header_footer(doc)

    lines = md_text.split("\n")
    i = 0
    in_table = False
    table_lines: list[str] = []
    in_blockquote = False
    blockquote_lines: list[str] = []
    metadata_lines: list[str] = []
    collecting_metadata = False
    first_heading_seen = False
    h1_count = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("<!--"):
            while "-->" not in stripped and i < len(lines) - 1:
                i += 1
                stripped = lines[i].strip()
            i += 1
            continue

        if in_blockquote and not stripped.startswith(">"):
            _flush_blockquote(doc, blockquote_lines)
            in_blockquote = False
            blockquote_lines = []

        if in_table and not (stripped.startswith("|") and "|" in stripped[1:]):
            _add_table(doc, table_lines)
            in_table = False
            table_lines = []

        if collecting_metadata and not _is_metadata_line(stripped):
            _add_metadata_block(doc, metadata_lines)
            collecting_metadata = False
            metadata_lines = []

        if re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="2" w:space="4" '
                f'w:color="{DIVIDER_HEX}"/>'
                f'</w:pBdr>'
            )
            p._p.get_or_add_pPr().append(pBdr)
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(stripped)
            i += 1
            continue

        if stripped.startswith(">"):
            in_blockquote = True
            bq_text = stripped[1:].strip()
            if bq_text:
                blockquote_lines.append(bq_text)
            i += 1
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.*)", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            if level == 1:
                h1_count += 1
                if h1_count > 1:
                    _add_page_break(doc)
            if not first_heading_seen:
                first_heading_seen = True
            h = doc.add_heading(heading_text, level=level)
            for run in h.runs:
                _set_run_font(run, size={1: 18, 2: 14, 3: 12}[level],
                              color=CHARCOAL, bold=True)
            if level == 2:
                _add_h2_underline(h)
            i += 1
            continue

        h4_match = re.match(r"^#{4,}\s+(.*)", stripped)
        if h4_match:
            p = doc.add_paragraph()
            run = p.add_run(h4_match.group(1).strip())
            _set_run_font(run, size=11, color=CHARCOAL, bold=True)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        if _is_metadata_line(stripped) and not first_heading_seen:
            collecting_metadata = True
            metadata_lines.append(stripped)
            i += 1
            continue
        if _is_metadata_line(stripped) and collecting_metadata:
            metadata_lines.append(stripped)
            i += 1
            continue

        indent_match = re.match(r"^(\s{2,})([-*])\s+(.*)", line)
        if indent_match:
            bullet_text = indent_match.group(3)
            try:
                p = doc.add_paragraph(style="List Bullet 2")
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Cm(2.0)
            _add_rich_text(p, bullet_text, base_size=10.5)
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            bullet_text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_rich_text(p, bullet_text)
            i += 1
            continue

        num_match = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if num_match:
            p = doc.add_paragraph(style="List Number")
            _add_rich_text(p, num_match.group(2))
            i += 1
            continue

        if _is_metadata_line(stripped) and first_heading_seen:
            m = _METADATA_RE.match(stripped)
            if m:
                p = doc.add_paragraph()
                run_key = p.add_run(f"{m.group(1)}: ")
                _set_run_font(run_key, size=11, color=BODY_TEXT, bold=True)
                run_val = p.add_run(m.group(2))
                _set_run_font(run_val, size=11, color=BODY_TEXT)
                i += 1
                continue

        if not stripped:
            i += 1
            continue

        p = doc.add_paragraph()
        _add_rich_text(p, stripped)
        i += 1

    if in_table and table_lines:
        _add_table(doc, table_lines)
    if in_blockquote and blockquote_lines:
        _flush_blockquote(doc, blockquote_lines)
    if collecting_metadata and metadata_lines:
        _add_metadata_block(doc, metadata_lines)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run(
        "Draft prepared by CDCN Agent — requires review and approval before use."
    )
    _set_run_font(run, size=8.5, color=SUBTLE_GREY, italic=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out_path = OUTPUT_DIR / filename
    doc.save(str(out_path))
    log.info("DOCX saved (legacy): %s", out_path)
    return out_path


# ── Main converter ───────────────────────────────────────────────────────────

def markdown_to_docx(md_text: str, filename: str = "", title: str = "") -> Path:
    """
    Convert markdown text to a professionally formatted DOCX file.

    Uses the Node.js docx_generator.js for superior formatting with full
    CDCN branding (title page, resolution boxes, etc.).  Falls back to the
    legacy python-docx renderer if the JS generator fails.

    Returns the Path to the saved file.
    """
    # Resolve filename
    if not filename:
        stamp = datetime.now().strftime("%Y%m%d_%H%M")
        filename = f"CDCN_Draft_{stamp}.docx"
    if not filename.endswith(".docx"):
        filename += ".docx"
    filename = re.sub(r'[<>:"/\\|?*]', '_', filename)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # ── Try JS generator (primary path) ──────────────────────────────────
    try:
        json_payload = _markdown_to_json(md_text, title=title)
        out_path = _call_js_generator(json_payload, filename)
        log.info("DOCX saved (JS): %s", out_path)
        return out_path
    except Exception as exc:
        log.warning("JS generator failed, falling back to python-docx: %s", exc)

    # ── Fallback to legacy python-docx ───────────────────────────────────
    return _legacy_markdown_to_docx(md_text, filename)


def get_download_url(filepath: Path) -> str:
    """Return the web-accessible download URL for a document."""
    import urllib.parse
    name = filepath.name
    encoded = urllib.parse.quote(name)
    return f"{DOWNLOAD_BASE}{encoded}"
