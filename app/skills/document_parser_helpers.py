"""
Document parsing helpers — PDF and DOCX extraction routines.

Extracted from document_parser.py to keep each module under 500 lines.
"""
import logging
import subprocess
import zipfile
from pathlib import Path

log = logging.getLogger(__name__)


# ── PDF parsing ──────────────────────────────────────────────────────────────

def _ocr_pdf(path: Path) -> list[str]:
    """OCR a scanned PDF using pytesseract.  Returns list of page texts."""
    pages: list[str] = []
    try:
        import pytesseract
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                try:
                    img = page.to_image(resolution=300).original  # PIL Image
                    from app.config import settings as _cfg
                    text = pytesseract.image_to_string(
                        img, timeout=_cfg.ocr_timeout_secs,
                    )
                    if text.strip():
                        pages.append(text.strip())
                except Exception as exc:
                    log.debug("OCR failed for page %d of %s: %s", i + 1, path.name, exc)
    except ImportError:
        log.warning("pytesseract not installed — OCR unavailable for %s", path.name)
    except Exception as exc:
        log.warning("OCR pipeline failed for %s: %s", path.name, exc)
    return pages


def _pdf_via_pdfplumber(path: Path) -> tuple[str, list, int]:
    """Attempt 1: extract text and tables via pdfplumber."""
    import pdfplumber
    from app.skills.document_parser import ParsedTable, _table_to_markdown

    all_text_parts: list[str] = []
    all_tables: list = []

    with pdfplumber.open(path) as pdf:
        page_count = len(pdf.pages)
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text() or ""
            if text:
                all_text_parts.append(text)

            tables = page.extract_tables()
            for table in tables:
                md = _table_to_markdown(table)
                if md:
                    all_tables.append(ParsedTable(markdown=md, page_number=i))

    full_text = "\n\n".join(all_text_parts)
    return full_text, all_tables, page_count


def _pdf_via_pdftotext(path: Path, timeout: int = 10) -> str:
    """Attempt 2: extract text via pdftotext subprocess."""
    try:
        result = subprocess.run(
            ["pdftotext", "-layout", str(path), "-"],
            capture_output=True,
            text=True,
            timeout=timeout,
        )
        if result.returncode != 0:
            raise RuntimeError(f"pdftotext exited with code {result.returncode}: {result.stderr.strip()}")
        return result.stdout.strip()
    except FileNotFoundError:
        raise RuntimeError("pdftotext binary not found on system")


def _pdf_via_ocr(path: Path) -> str:
    """Attempt 3: extract text via pytesseract OCR."""
    ocr_pages = _ocr_pdf(path)
    if not ocr_pages:
        raise RuntimeError("OCR produced no text")
    return "\n\n".join(ocr_pages)


def _parse_pdf(path: Path) -> "ParsedDocument":
    """Parse a PDF using a fallback chain: pdfplumber -> pdftotext -> OCR."""
    from app.skills.document_parser import (
        ParsedDocument, _extract_title, _extract_date_from_text,
        _extract_attendees, _extract_apologies, _extract_chair,
        _detect_sections,
    )

    extraction_method = ""
    full_text = ""
    all_tables = []
    page_count = 0
    failure_reasons: list[str] = []

    # Attempt 1: pdfplumber
    try:
        log.info("PDF parse attempt 1 (pdfplumber) for %s", path.name)
        full_text, all_tables, page_count = _pdf_via_pdfplumber(path)
        chars_per_page = len(full_text.strip()) / max(page_count, 1)
        if chars_per_page < 100:
            reason = f"pdfplumber low text density ({chars_per_page:.0f} chars/page)"
            log.info("%s for %s — trying next method", reason, path.name)
            failure_reasons.append(reason)
            full_text = ""  # trigger fallback
        else:
            extraction_method = "pdfplumber"
    except Exception as exc:
        reason = f"pdfplumber failed: {type(exc).__name__}: {exc}"
        log.warning("%s for %s", reason, path.name)
        failure_reasons.append(reason)

    # Attempt 2: pdftotext
    if not full_text:
        try:
            log.info(
                "PDF parse attempt 2 (pdftotext) for %s (reason: %s)",
                path.name, failure_reasons[-1] if failure_reasons else "unknown",
            )
            full_text = _pdf_via_pdftotext(path, timeout=10)
            if full_text:
                extraction_method = "pdftotext"
            else:
                reason = "pdftotext returned empty text"
                log.info("%s for %s", reason, path.name)
                failure_reasons.append(reason)
                full_text = ""
        except (RuntimeError, subprocess.TimeoutExpired, OSError) as exc:
            reason = f"pdftotext failed: {type(exc).__name__}: {exc}"
            log.warning("%s for %s", reason, path.name)
            failure_reasons.append(reason)

    # Attempt 3: OCR
    if not full_text:
        try:
            log.info(
                "PDF parse attempt 3 (OCR) for %s (reason: %s)",
                path.name, failure_reasons[-1] if failure_reasons else "unknown",
            )
            full_text = _pdf_via_ocr(path)
            extraction_method = "ocr"
            log.info("OCR extracted %d chars from %s", len(full_text), path.name)
        except (RuntimeError, ImportError, OSError) as exc:
            reason = f"OCR failed: {type(exc).__name__}: {exc}"
            log.warning("%s for %s", reason, path.name)
            failure_reasons.append(reason)

    # All methods failed
    if not full_text:
        error_detail = "; ".join(failure_reasons)
        log.error("All PDF extraction methods failed for %s: %s", path.name, error_detail)
        return ParsedDocument(
            full_text="",
            title=path.stem,
            page_count=page_count,
            extraction_method="none",
        )

    title = _extract_title(full_text, path.name)
    date = _extract_date_from_text(full_text)
    attendees = _extract_attendees(full_text)
    apologies = _extract_apologies(full_text)
    chair = _extract_chair(full_text)
    sections = _detect_sections(full_text)

    return ParsedDocument(
        full_text=full_text,
        sections=sections,
        tables=all_tables,
        title=title,
        date=date,
        attendees=attendees,
        page_count=page_count,
        extraction_method=extraction_method,
        apologies=apologies,
        chair=chair,
    )


# ── DOCX parsing ─────────────────────────────────────────────────────────────

def _parse_docx(path: Path) -> "ParsedDocument":
    """Parse a DOCX using python-docx, detecting headings and bold text."""
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
    from app.skills.document_parser import (
        ParsedDocument, ParsedSection, ParsedTable,
        _is_section_heading, _detect_sections,
        _extract_title, _extract_date_from_text,
        _extract_attendees, _extract_apologies, _extract_chair,
        _table_to_markdown,
    )

    # Check file size and warn for large files
    file_size = path.stat().st_size
    if file_size > 10 * 1024 * 1024:
        log.warning(
            "Large DOCX file (%d MB): %s — processing may be slow",
            file_size // (1024 * 1024), path.name,
        )

    # Handle corrupted zip / password-protected files
    try:
        doc = Document(str(path))
    except zipfile.BadZipFile:
        log.error("Corrupted DOCX (bad zip): %s", path.name)
        return ParsedDocument(
            full_text="",
            title=path.stem,
            extraction_method="none",
        )
    except PackageNotFoundError as exc:
        log.error("Cannot open DOCX package for %s: %s", path.name, exc)
        return ParsedDocument(
            full_text="",
            title=path.stem,
            extraction_method="none",
        )
    except KeyError as exc:
        log.error("Possible password-protected DOCX: %s (%s)", path.name, exc)
        return ParsedDocument(
            full_text="",
            title=path.stem,
            extraction_method="none",
        )

    all_lines: list[str] = []
    sections: list[ParsedSection] = []
    current_title = "Preamble"
    current_lines: list[str] = []
    has_images = False

    # Check for embedded images in the document
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            has_images = True
            break

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            all_lines.append("")
            current_lines.append("")
            continue

        # Detect headings: actual heading style, or bold + short
        is_heading = False
        if para.style and para.style.name.startswith("Heading"):
            is_heading = True
        elif len(text) < 120 and para.runs and all(r.bold for r in para.runs if r.text.strip()):
            is_heading = True
        elif _is_section_heading(text):
            is_heading = True

        if is_heading:
            content = "\n".join(current_lines).strip()
            if content:
                sections.append(ParsedSection(title=current_title, content=content))
            current_title = text
            current_lines = []
        else:
            current_lines.append(text)

        all_lines.append(text)

    # Save last section
    content = "\n".join(current_lines).strip()
    if content:
        sections.append(ParsedSection(title=current_title, content=content))

    # Extract tables from DOCX
    all_tables: list = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        md = _table_to_markdown(rows)
        if md:
            all_tables.append(ParsedTable(markdown=md))

    full_text = "\n".join(all_lines)

    # Warn if document has images but no text (likely scanned)
    if has_images and len(full_text.strip()) < 50:
        log.warning("DOCX contains images but no meaningful text (possible scanned document): %s", path.name)

    title = _extract_title(full_text, path.name)
    date = _extract_date_from_text(full_text)
    attendees = _extract_attendees(full_text)
    apologies = _extract_apologies(full_text)
    chair = _extract_chair(full_text)

    return ParsedDocument(
        full_text=full_text,
        sections=sections if len(sections) > 1 else _detect_sections(full_text),
        tables=all_tables,
        title=title,
        date=date,
        attendees=attendees,
        page_count=0,
        extraction_method="python-docx",
        has_images=has_images,
        apologies=apologies,
        chair=chair,
    )
