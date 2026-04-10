"""
Helpers for the DOCX converter — brand constants, font/style setup, and
low-level DOCX element builders.

Extracted from docx_converter.py to keep each module under 500 lines.
"""
import re
import logging

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

log = logging.getLogger(__name__)

# ── Brand colours ─────────────────────────────────────────────────────────────
FONT_FAMILY = "Rubik"
FONT_FALLBACK = "Calibri"
CODE_FONT = "Consolas"

CHARCOAL = RGBColor(0x2C, 0x3E, 0x50)        # #2C3E50 — headings, table header bg
SAILING_BLUE = RGBColor(0x5E, 0xAF, 0xE5)     # #5EAFE5 — accent, H2 underline
BODY_TEXT = RGBColor(0x33, 0x33, 0x33)         # #333333 — body text
SUBTLE_GREY = RGBColor(0x94, 0x94, 0x94)       # #949494 — footer, light text
META_GREY = RGBColor(0x5A, 0x5A, 0x5A)         # #5A5A5A — metadata labels
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Hex versions for XML shading (no #)
CHARCOAL_HEX = "2C3E50"
SAILING_BLUE_HEX = "5EAFE5"
TABLE_ALT_HEX = "F2F5F8"
BORDER_GREY_HEX = "D0D0D0"
DIVIDER_HEX = "D0D0D0"


# ── Font helper ──────────────────────────────────────────────────────────────

def _set_run_font(run, size: float = 11, color: RGBColor = BODY_TEXT,
                  bold: bool = False, italic: bool = False,
                  font: str = FONT_FAMILY) -> None:
    """Configure a run with the brand font and styling."""
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    # Set east-asia / complex-script fallback so Word doesn't override
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), FONT_FALLBACK)


# ── Style setup ──────────────────────────────────────────────────────────────

def _setup_styles(doc: Document) -> None:
    """Define or update the core Word styles with CDCN branding."""
    styles = doc.styles

    # ── Normal ────────────────────────────────────────────────────────────
    normal = styles["Normal"]
    normal.font.name = FONT_FAMILY
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY_TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    # Set font in the style XML so it actually sticks
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), FONT_FAMILY)
    rFonts.set(qn("w:hAnsi"), FONT_FAMILY)
    rFonts.set(qn("w:cs"), FONT_FALLBACK)

    # ── Heading 1 — 18pt, Charcoal, bold, 24pt space before ─────────────
    h1 = styles["Heading 1"]
    h1.font.name = FONT_FAMILY
    h1.font.size = Pt(18)
    h1.font.color.rgb = CHARCOAL
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    # ── Heading 2 — 14pt, Charcoal, bold ─────────────────────────────────
    h2 = styles["Heading 2"]
    h2.font.name = FONT_FAMILY
    h2.font.size = Pt(14)
    h2.font.color.rgb = CHARCOAL
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.keep_with_next = True

    # ── Heading 3 — 12pt, Charcoal, bold ─────────────────────────────────
    h3 = styles["Heading 3"]
    h3.font.name = FONT_FAMILY
    h3.font.size = Pt(12)
    h3.font.color.rgb = CHARCOAL
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.keep_with_next = True


# ── Cell / page helpers ──────────────────────────────────────────────────────

def _set_cell_shading(cell, hex_color: str) -> None:
    """Apply background shading to a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _add_page_number(paragraph) -> None:
    """Insert a PAGE field into a paragraph (for footer)."""
    run = paragraph.add_run()
    fld_char_begin = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
    )
    run._r.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
    )
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fld_char_end = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
    )
    run3._r.append(fld_char_end)


def _add_header_footer(doc: Document) -> None:
    """Add branded header and page-number footer to all sections."""
    for section in doc.sections:
        # ── Header ────────────────────────────────────────────────────
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.clear()
        run = hp.add_run("Community Development Company of Nesting")
        _set_run_font(run, size=8, color=SUBTLE_GREY)
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.paragraph_format.space_after = Pt(0)

        # Thin blue line under header
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="4" '
            f'w:color="{SAILING_BLUE_HEX}"/>'
            f'</w:pBdr>'
        )
        hp._p.get_or_add_pPr().append(pBdr)

        # ── Footer ────────────────────────────────────────────────────
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        run = fp.add_run("CDCN  ·  SC048164  ·  Page ")
        _set_run_font(run, size=7.5, color=SUBTLE_GREY)
        _add_page_number(fp)
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)


def _add_h2_underline(paragraph) -> None:
    """Add a Sailing Blue underline beneath a Heading 2 paragraph."""
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="3" '
        f'w:color="{SAILING_BLUE_HEX}"/>'
        f'</w:pBdr>'
    )
    paragraph._p.get_or_add_pPr().append(pBdr)


def _add_page_break(doc: Document) -> None:
    """Insert a page break."""
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


# ── Inline markdown parsing ──────────────────────────────────────────────────

# Order matters: bold-italic before bold before italic
_INLINE_PATTERN = re.compile(
    r"(\*\*\*.*?\*\*\*)"   # ***bold italic***
    r"|(\*\*.*?\*\*)"       # **bold**
    r"|(\*[^*]+?\*)"        # *italic*
    r"|(`[^`]+?`)"          # `code`
)


def _add_rich_text(paragraph, text: str, base_size: float = 11,
                   base_color: RGBColor = BODY_TEXT,
                   base_bold: bool = False) -> None:
    """Parse inline markdown and add formatted runs to a paragraph."""
    if not text:
        return

    last_end = 0
    for m in _INLINE_PATTERN.finditer(text):
        # Add plain text before this match
        if m.start() > last_end:
            run = paragraph.add_run(text[last_end:m.start()])
            _set_run_font(run, size=base_size, color=base_color, bold=base_bold)

        matched = m.group()
        if matched.startswith("***") and matched.endswith("***"):
            run = paragraph.add_run(matched[3:-3])
            _set_run_font(run, size=base_size, color=base_color,
                          bold=True, italic=True)
        elif matched.startswith("**") and matched.endswith("**"):
            run = paragraph.add_run(matched[2:-2])
            _set_run_font(run, size=base_size, color=base_color, bold=True)
        elif matched.startswith("*") and matched.endswith("*"):
            run = paragraph.add_run(matched[1:-1])
            _set_run_font(run, size=base_size, color=base_color, italic=True)
        elif matched.startswith("`") and matched.endswith("`"):
            run = paragraph.add_run(matched[1:-1])
            _set_run_font(run, size=base_size - 0.5, color=SAILING_BLUE,
                          font=CODE_FONT)
            last_end = m.end()
            continue

        last_end = m.end()

    # Add remaining plain text
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        _set_run_font(run, size=base_size, color=base_color, bold=base_bold)


# ── Table parsing ────────────────────────────────────────────────────────────

def _parse_table(lines: list[str]) -> list[list[str]]:
    """Parse markdown table lines into rows of cells."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        if re.match(r"^\|[\s\-:|]+\|$", line):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if cells:
            rows.append(cells)
    return rows


def _add_table(doc: Document, table_lines: list[str]) -> None:
    """Render a markdown table as a branded DOCX table."""
    rows = _parse_table(table_lines)
    if not rows:
        return

    n_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True

    # Table borders — clean horizontal lines, subtle verticals
    tbl_pr = tbl._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="6" w:space="0" '
        f'w:color="{CHARCOAL_HEX}"/>'
        f'  <w:left w:val="single" w:sz="2" w:space="0" '
        f'w:color="{BORDER_GREY_HEX}"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" '
        f'w:color="{CHARCOAL_HEX}"/>'
        f'  <w:right w:val="single" w:sz="2" w:space="0" '
        f'w:color="{BORDER_GREY_HEX}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" '
        f'w:color="{BORDER_GREY_HEX}"/>'
        f'  <w:insideV w:val="single" w:sz="2" w:space="0" '
        f'w:color="{BORDER_GREY_HEX}"/>'
        f'</w:tblBorders>'
    )
    tbl_pr.append(borders)

    # Cell margins for breathing room
    tbl_margins = parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        f'  <w:top w:w="60" w:type="dxa"/>'
        f'  <w:left w:w="100" w:type="dxa"/>'
        f'  <w:bottom w:w="60" w:type="dxa"/>'
        f'  <w:right w:w="100" w:type="dxa"/>'
        f'</w:tblCellMar>'
    )
    tbl_pr.append(tbl_margins)

    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci >= n_cols:
                break
            cell = tbl.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]

            if ri == 0:
                # Header row: Charcoal background, white bold text
                _set_cell_shading(cell, CHARCOAL_HEX)
                _add_rich_text(p, cell_text.replace("**", ""),
                               base_size=10, base_color=WHITE, base_bold=True)
            else:
                # Alternating row shading
                if ri % 2 == 0:
                    _set_cell_shading(cell, TABLE_ALT_HEX)
                _add_rich_text(p, cell_text.replace("**", ""), base_size=10)

            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)


# ── Metadata block ──────────────────────────────────────────────────────────

_METADATA_RE = re.compile(r"^\*\*(.+?):\*\*\s*(.+)$")


def _is_metadata_line(line: str) -> bool:
    """Check if a line is a **Key:** Value metadata pair."""
    return bool(_METADATA_RE.match(line.strip()))


def _add_metadata_block(doc: Document, lines: list[str]) -> None:
    """Render a block of **Key:** Value lines as a styled metadata section."""
    for line in lines:
        m = _METADATA_RE.match(line.strip())
        if not m:
            continue
        key, value = m.group(1), m.group(2)
        p = doc.add_paragraph()
        run_key = p.add_run(f"{key}: ")
        _set_run_font(run_key, size=10, color=META_GREY, bold=True)
        run_val = p.add_run(value)
        _set_run_font(run_val, size=10, color=BODY_TEXT)


def _flush_blockquote(doc: Document, blockquote_lines: list[str]) -> None:
    """Render accumulated blockquote lines as a styled callout box."""
    for bq_line in blockquote_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        # Blue left border
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:left w:val="single" w:sz="12" w:space="8" '
            f'w:color="{SAILING_BLUE_HEX}"/>'
            f'</w:pBdr>'
        )
        p._p.get_or_add_pPr().append(pBdr)
        _add_rich_text(p, bq_line, base_size=10.5, base_color=META_GREY)
